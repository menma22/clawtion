"""Multi-format file loaders for clawtion indexer.

Provides FileProcessor implementations for HTML, CSV, JSON, EPUB,
and DOCX files.  Optional dependencies (beautifulsoup4, ebooklib,
python-docx) are imported lazily so the application works without them.
"""

from __future__ import annotations

import contextlib
import csv
import json
import os
from html.parser import HTMLParser
from typing import Any, ClassVar

from clawtion.utils.logging import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Optional dependency detection
# ---------------------------------------------------------------------------

_BS4_AVAILABLE: bool = False
try:
    import bs4  # noqa: F401

    _BS4_AVAILABLE = True
except ImportError:
    pass

_EBOOKLIB_AVAILABLE: bool = False
try:
    import ebooklib  # noqa: F401

    _EBOOKLIB_AVAILABLE = True
except ImportError:
    pass

_DOCX_AVAILABLE: bool = False
try:
    import docx  # noqa: F401

    _DOCX_AVAILABLE = True
except ImportError:
    pass


# ---------------------------------------------------------------------------
# HTML Processor
# ---------------------------------------------------------------------------


class _HTMLTextExtractor(HTMLParser):
    """Simple HTML-to-text extractor using Python's stdlib HTMLParser.

    Strips all tags and decodes HTML entities.  Used when
    beautifulsoup4 is not installed.
    """

    def __init__(self) -> None:
        super().__init__()
        self._text_parts: list[str] = []
        self._skip: bool = False
        self._skip_tags: set[str] = {"script", "style", "noscript"}

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in self._skip_tags:
            self._skip = True
        if tag in ("br", "p", "tr", "li", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._text_parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._skip_tags:
            self._skip = False

    def handle_data(self, data: str) -> None:
        if not self._skip:
            self._text_parts.append(data)

    def get_text(self) -> str:
        return "".join(self._text_parts).strip()


class HTMLProcessor:
    """Extract text content from HTML files.

    Tries beautifulsoup4 first when available; falls back to stdlib
    :class:`html.parser.HTMLParser`.
    """

    def can_process(self, file_path: str) -> bool:
        ext = os.path.splitext(file_path)[1].lower()
        return ext in (".html", ".htm", ".xhtml")

    def extract_content(self, file_path: str) -> dict[str, Any]:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            raw = f.read()

        if _BS4_AVAILABLE:
            return self._extract_with_bs4(raw)

        return self._extract_with_stdlib(raw)

    def _extract_with_bs4(self, html: str) -> dict[str, Any]:
        import bs4

        soup = bs4.BeautifulSoup(html, "html.parser")
        # Remove script/style elements
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        # Collapse multiple newlines
        import re

        text = re.sub(r"\n{3,}", "\n\n", text)

        title = soup.title.string.strip() if soup.title and soup.title.string else ""

        return {
            "text": text,
            "metadata": {
                "title": title or None,
                "parser": "beautifulsoup4",
            },
        }

    def _extract_with_stdlib(self, html: str) -> dict[str, Any]:
        extractor = _HTMLTextExtractor()
        extractor.feed(html)
        text = extractor.get_text()

        # Try to extract title from <title> tag the simple way
        title = ""
        import re

        m = re.search(r"<title[^>]*>([^<]+)</title>", html, re.IGNORECASE)
        if m:
            title = m.group(1).strip()

        return {
            "text": text,
            "metadata": {
                "title": title or None,
                "parser": "stdlib.html.parser",
            },
        }

    def get_supported_extensions(self) -> list[str]:
        return [".html", ".htm", ".xhtml"]


# ---------------------------------------------------------------------------
# CSV Processor
# ---------------------------------------------------------------------------


class CSVProcessor:
    """Extract text from CSV files.

    Reads all rows and joins each row's values into a pipe-delimited line.
    """

    def can_process(self, file_path: str) -> bool:
        return os.path.splitext(file_path)[1].lower() == ".csv"

    def extract_content(self, file_path: str) -> dict[str, Any]:
        rows: list[dict[str, str]] = []
        fieldnames: list[str] = []
        total_rows: int = 0

        with open(file_path, encoding="utf-8", errors="replace", newline="") as f:
            sample = f.read(4096)
            f.seek(0)

            try:
                dialect = csv.Sniffer().sniff(sample)
            except csv.Error:
                dialect = csv.excel

            reader = csv.DictReader(f, dialect=dialect)
            if reader.fieldnames:
                fieldnames = list(reader.fieldnames)

            for row in reader:
                # Normalise all values to strings
                cleaned: dict[str, str] = {}
                for k, v in row.items():
                    cleaned[k or ""] = (v or "").strip()
                rows.append(cleaned)
                total_rows += 1

        # Build a human-readable text representation
        lines: list[str] = []

        if fieldnames:
            lines.append(" | ".join(fieldnames))
            lines.append("-" * len(" | ".join(fieldnames)))

        for row in rows:
            vals = [row.get(h, "") for h in fieldnames]
            lines.append(" | ".join(vals))

        text = "\n".join(lines)

        return {
            "text": text,
            "metadata": {
                "columns": fieldnames,
                "total_rows": total_rows,
                "parser": "csv.DictReader",
            },
        }

    def get_supported_extensions(self) -> list[str]:
        return [".csv"]


# ---------------------------------------------------------------------------
# JSON Processor
# ---------------------------------------------------------------------------


class JSONProcessor:
    """Extract text from JSON files.

    Flattens nested structures using dot-separated keys so the content
    is suitable for text search and embedding.
    """

    def can_process(self, file_path: str) -> bool:
        return os.path.splitext(file_path)[1].lower() == ".json"

    def extract_content(self, file_path: str) -> dict[str, Any]:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            data = json.load(f)

        lines = self._flatten(data, prefix="")
        text = "\n".join(lines)

        return {
            "text": text,
            "metadata": {
                "root_type": type(data).__name__,
                "parser": "json.flatten",
            },
        }

    def _flatten(self, obj: Any, prefix: str = "") -> list[str]:
        """Recursively flatten a JSON object into ``key: value`` lines."""
        parts: list[str] = []

        if isinstance(obj, dict):
            for key, value in obj.items():
                full_key = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    child = self._flatten(value, full_key)
                    if child:
                        parts.extend(child)
                else:
                    parts.append(f"{full_key}: {self._format_value(value)}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                idx_key = f"{prefix}[{i}]"
                if isinstance(item, (dict, list)):
                    child = self._flatten(item, idx_key)
                    if child:
                        parts.extend(child)
                else:
                    parts.append(f"{idx_key}: {self._format_value(item)}")
        else:
            parts.append(f"{prefix}: {self._format_value(obj)}")

        return parts

    @staticmethod
    def _format_value(value: Any) -> str:
        if value is None:
            return "null"
        if isinstance(value, bool):
            return str(value).lower()
        return str(value)

    def get_supported_extensions(self) -> list[str]:
        return [".json"]


# ---------------------------------------------------------------------------
# EPUB Processor
# ---------------------------------------------------------------------------


class EPUBProcessor:
    """Extract text from EPUB e-book files.

    Requires the ``ebooklib`` and ``BeautifulSoup`` packages::

        pip install ebooklib beautifulsoup4

    Gracefully falls back with a descriptive message when the libraries
    are not installed.
    """

    _DEPENDENCY_HINT: ClassVar[str] = (
        "Install optional dependencies: pip install clawtion[epub]  (requires ebooklib and beautifulsoup4)"
    )

    def can_process(self, file_path: str) -> bool:
        return os.path.splitext(file_path)[1].lower() == ".epub"

    def extract_content(self, file_path: str) -> dict[str, Any]:
        if not _EBOOKLIB_AVAILABLE:
            logger.warning("ebooklib is not installed. Cannot process EPUB file.", path=file_path)
            return {
                "text": "",
                "metadata": {
                    "error": "Missing dependency: ebooklib",
                    "hint": self._DEPENDENCY_HINT,
                },
            }

        import ebooklib
        from ebooklib import epub

        book = epub.read_epub(file_path)

        # Build a flat list of document items (skip images / styles)
        doc_items: list[Any] = []
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                doc_items.append(item)

        chapters: list[str] = []
        chapter_titles: list[str] = []

        for item in doc_items:
            content_bytes = item.get_body_content()
            html_content = content_bytes.decode("utf-8", errors="replace")

            if _BS4_AVAILABLE:
                import bs4

                soup = bs4.BeautifulSoup(html_content, "html.parser")
                title_tag = soup.find("title")
                if title_tag and title_tag.string:
                    chapter_titles.append(title_tag.string.strip())

                for tag in soup(["script", "style", "noscript"]):
                    tag.decompose()

                text = soup.get_text(separator="\n", strip=True)
            else:
                # Use the stdlib extractor
                extractor = _HTMLTextExtractor()
                extractor.feed(html_content)
                text = extractor.get_text()

            if text.strip():
                chapters.append(text.strip())

        text = "\n\n".join(chapters)

        # Extract metadata from OPF
        meta_title: str = ""
        meta_author: str = ""
        with contextlib.suppress(IndexError, TypeError):
            meta_title = book.get_metadata("DC", "title")[0][0]  # type: ignore[arg-type]
        with contextlib.suppress(IndexError, TypeError):
            meta_author = book.get_metadata("DC", "creator")[0][0]  # type: ignore[arg-type]

        return {
            "text": text,
            "metadata": {
                "title": meta_title or None,
                "author": meta_author or None,
                "chapters": len(chapters),
                "parser": "ebooklib",
            },
        }

    def get_supported_extensions(self) -> list[str]:
        return [".epub"]


# ---------------------------------------------------------------------------
# DOCX Processor
# ---------------------------------------------------------------------------


class DocxProcessor:
    """Extract text from Word ``.docx`` files.

    Requires the ``python-docx`` package::

        pip install python-docx

    Gracefully falls back with a descriptive message when the library
    is not installed.
    """

    _DEPENDENCY_HINT: ClassVar[str] = "Install optional dependency: pip install python-docx"

    def can_process(self, file_path: str) -> bool:
        return os.path.splitext(file_path)[1].lower() == ".docx"

    def extract_content(self, file_path: str) -> dict[str, Any]:
        if not _DOCX_AVAILABLE:
            logger.warning("python-docx is not installed. Cannot process DOCX file.", path=file_path)
            return {
                "text": "",
                "metadata": {
                    "error": "Missing dependency: python-docx",
                    "hint": self._DEPENDENCY_HINT,
                },
            }

        import docx

        document = docx.Document(file_path)

        paragraphs: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        text = "\n\n".join(paragraphs)

        # Extract metadata from core properties if available
        metadata: dict[str, Any] = {
            "parser": "python-docx",
            "paragraph_count": len(paragraphs),
        }

        try:
            core_props = document.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
        except Exception:
            pass

        # Also extract text from tables
        table_texts: list[str] = []
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))

        if table_texts:
            text += "\n\n--- Tables ---\n\n" + "\n\n".join(table_texts)
            metadata["table_count"] = len(document.tables)

        return {
            "text": text,
            "metadata": metadata,
        }

    def get_supported_extensions(self) -> list[str]:
        return [".docx"]
